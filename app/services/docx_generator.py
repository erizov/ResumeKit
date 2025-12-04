"""
DOCX generation service for tailored resumes.

Converts resume text to DOCX format using python-docx.
"""

import io
import re

try:
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    Document = None  # type: ignore


def generate_docx_from_text(resume_text: str) -> bytes:
    """
    Generate a DOCX file from resume text.

    Converts plain text resume to DOCX with basic formatting.

    Args:
        resume_text: Plain text content of the resume.

    Returns:
        DOCX file content as bytes.

    Raises:
        RuntimeError: If python-docx is not installed or DOCX generation fails.
    """
    if Document is None:
        raise RuntimeError(
            "python-docx is not installed. Install it with: pip install python-docx"
        )

    # Clean the text before processing
    from .resume_formatter import clean_resume_text
    resume_text = clean_resume_text(resume_text)

    doc = Document()
    
    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Parse text into paragraphs
    lines = resume_text.splitlines()
    prev_empty = False
    
    for line in lines:
        line = line.strip()
        if not line:
            # Only add one empty paragraph for spacing
            if not prev_empty:
                doc.add_paragraph()
                prev_empty = True
            continue
        
        prev_empty = False

        # Detect headers (all caps or lines ending with colon)
        if line.isupper() or line.endswith(":"):
            p = doc.add_paragraph(line)
            p.style = "Heading 2"
            p.runs[0].font.size = Pt(14)
            p.runs[0].bold = True
        # Detect bullet points
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
        # Check if line has date pattern (job title with date)
        elif re.search(r'\d{4}\s*[–-]\s*\d{4}', line):
            # Format with date on right using tab
            parts = re.split(r'(\d{4}\s*[–-]\s*\d{4})', line)
            if len(parts) >= 2:
                title_part = parts[0].strip()
                date_part = parts[1]
                p = doc.add_paragraph()
                p.add_run(title_part)
                p.add_run("\t" * 8)  # Tab to push date to right
                p.add_run(date_part).bold = True
            else:
                doc.add_paragraph(line)
        else:
            doc.add_paragraph(line)

    # Save to bytes buffer
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer.getvalue()

